# -*- coding: utf-8 -*-
"""生成《AITC 对外 HTTP 接口说明》Word 文档。

用法：C:\\Users\\Finn\\.conda\\envs\\aitc\\python.exe gen_api_docs.py
输出：docs/对外HTTP接口说明.docx
"""

from __future__ import annotations

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "docs")
OUT_FILE = os.path.join(OUT_DIR, "对外HTTP接口说明.docx")

BASE = "http://127.0.0.1:8088"

# ---------------------------------------------------------------------------
# 接口数据
# ---------------------------------------------------------------------------
GENERAL = """统一约定：
- 服务地址：{base}（HTTP 服务端口 8088，TCP 数据端口 65432）
- 请求/响应均为 JSON（Content-Type: application/json），编码 UTF-8
- 成功：HTTP 200，按接口返回对应结构；参数错误 HTTP 400；服务异常 HTTP 500
- 跨域：拒绝跨源请求（需同源或经网关转发）
- 绿波接口支持文档1（segment_id/green_wave_info）与文档2（corridor）两种请求格式，响应统一为文档1格式""".format(base=BASE)

INTERFACES = [
    # ---------------- 基础 ----------------
    {
        "name": "健康检查",
        "method": "GET",
        "path": "/health",
        "params": "无",
        "desc": "检查服务运行状态与数据质量",
        "req": "无",
        "resp": '{"status":"running","service":"aitc_runtime_http","radar_cache_size":0,"data_quality":{"total_issues":0,"issues_by_kind":{},"recent_issues":[]}}',
    },
    {
        "name": "前端页面",
        "method": "GET",
        "path": "/ 或 /index.html",
        "params": "无",
        "desc": "返回单文件前端页面（路口方案 / 放行控制流程 / 接口测试）",
        "req": "无",
        "resp": "HTML 页面",
    },
    # ---------------- 道路/路口配置 ----------------
    {
        "name": "查询道路信息",
        "method": "GET",
        "path": "/road_info/{cross_id}",
        "params": "cross_id：路口编号（数字字符串）",
        "desc": "查询指定路口道路配置（相位、最大/最小通行时间、平台限值）",
        "req": "无",
        "resp": '{"0":{"phase":["UD1","UD2","UDL","LR1","LR2","LRL",0,0,0,0],"max_pass_time":[35,55,64,35,56,64,0,0,0,0],...}}',
    },
    {
        "name": "查询路口信息",
        "method": "GET",
        "path": "/cross_info/{cross_id}",
        "params": "cross_id：路口编号（数字字符串）",
        "desc": "查询指定路口信息（相位映射、车道、路口编号）",
        "req": "无",
        "resp": '{"phase":{"1":"UD","2":"UD","3":"LR","4":"LRL",...},"LaneNo":{...},"jtll_ddbh":{...}}',
    },
    {
        "name": "新增/更新道路信息",
        "method": "POST",
        "path": "/road_info/update",
        "params": "body：道路配置对象（含 cross_id 等完整字段）",
        "desc": "新增或更新道路配置并写入 lib/road_info.json",
        "req": '{"Cross_id":"1300067","road_info":{...}}',
        "resp": '{"status":"success","saved":true,"operation":"updated","Cross_id":"1300067"}',
    },
    {
        "name": "新增/更新路口信息",
        "method": "POST",
        "path": "/cross_info/update",
        "params": "body：路口配置对象",
        "desc": "新增或更新路口配置并写入 lib/cross_info.json",
        "req": '{"Cross_id":"1300067","cross_info":{...}}',
        "resp": '{"status":"success","saved":true,"operation":"updated","Cross_id":"1300067"}',
    },
    # ---------------- 方案/Agent ----------------
    {
        "name": "单路口方案（直连DQN）",
        "method": "POST",
        "path": "/api/signal-timing",
        "params": "body.cross_id：路口编号",
        "desc": "直接调用单路口放行时间工具（DQN），不经大模型",
        "req": '{"cross_id":"1300068"}',
        "resp": '{"status":"success","cross_id":"1300068","result":{"cross_id":"1300068","signal_timing":[...],"coordinate_set":{"s1":0,"s2":0},"model_info":[...]}}',
    },
    {
        "name": "Qwen 编排方案",
        "method": "POST",
        "path": "/api/agent/signal-timing",
        "params": "body.cross_id；body.request_text：自然语言需求",
        "desc": "Qwen 先选工具，再调用单路口方案工具，返回中文汇总",
        "req": '{"cross_id":"1300068","request_text":"请给出当前这个路口的放行方案"}',
        "resp": '{"status":"success","cross_id":"1300068","result":{"status":"ok","summary":"已生成路口 1300068 的放行时间。","data":{"action":"signal.timing.single",...},"meta":{"llm_model":"deepseek-v4-flash"}}}',
    },
    {
        "name": "分步放行控制",
        "method": "POST",
        "path": "/api/agent/control-process",
        "params": "body.cross_id；body.request_text（可选）",
        "desc": "10 步规则函数判断 + 大模型逐步思考，返回每步思考与规则结果",
        "req": '{"cross_id":"1300068","request_text":"请分析该路口的放行方案"}',
        "resp": '{"status":"success","cross_id":"1300068","result":{"status":"ok","data":{"data_source":"data_hub","steps":[{step,key,title,llm_thought,data},...]},"meta":{"step_count":10}}}',
    },
    {
        "name": "雷达数据接收",
        "method": "POST",
        "path": "/",
        "params": "body：雷达/视频感知数据（单条或数组）",
        "desc": "HTTP 数据入口，按字段分类（deviceNo→radar 等）存入数据底座",
        "req": '{"deviceNo":"DEV001","eventType":null,...} 或 [...]',
        "resp": '{"status":"success","message":"Radar data received"}',
    },
    # ---------------- 绿波查询 ----------------
    {
        "name": "绿波-运行状态",
        "method": "GET",
        "path": "/api/green-wave/status",
        "params": "无",
        "desc": "查询绿波协调当前运行状态（启用/周期/相位差阶段/方向/走廊）",
        "req": "无",
        "resp": '{"status":"success","running":false,"cycle_green":false,"offset_green":false,"cycle_done":false,"offset_done":false,"round":0,"direction":"","corridor_id":"lvbo_01","last_send_time":0,"has_fix_plan":false}',
    },
    {
        "name": "绿波-走廊列表",
        "method": "GET",
        "path": "/api/green-wave/config",
        "params": "无",
        "desc": "查询全部绿波走廊完整配置",
        "req": "无",
        "resp": '{"status":"success","saved":false,"operation":"listed","items":[{corridor_id,name,enabled,cycle_seconds,intersections,periods,...}]}',
    },
    {
        "name": "绿波-单条走廊",
        "method": "GET",
        "path": "/api/green-wave/config/{corridor_id}",
        "params": "corridor_id：走廊编号（如 lvbo_01）",
        "desc": "查询指定走廊完整配置",
        "req": "无",
        "resp": '{"status":"success","saved":false,"operation":"queried","corridor":{corridor_id,name,enabled,cycle_seconds,intersections,periods,road_junction,topology}}',
    },
    {
        "name": "绿波-最新方案",
        "method": "GET",
        "path": "/api/green-wave/plan",
        "params": "无",
        "desc": "查询最新一轮实际下发的绿波方案",
        "req": "无",
        "resp": '{"status":"success","corridor_id":"lvbo_01","plan":{}}',
    },
    # ---------------- 绿波走廊 CRUD（文档1/文档2） ----------------
    {
        "name": "绿波走廊-列表（文档1）",
        "method": "GET",
        "path": "/green_wave?full=false|true",
        "params": "full：false 只返回编号/名称/启用；true 返回完整配置",
        "desc": "列出全部绿波走廊",
        "req": "无",
        "resp": '{"status":"success","saved":false,"operation":"listed","items":[{corridor_id,segment_id,name,enabled},...]}',
    },
    {
        "name": "绿波走廊-查询单条（文档1）",
        "method": "GET",
        "path": "/green_wave/{segment_id}",
        "params": "segment_id：走廊编号（如 lvbo_01）",
        "desc": "查询指定走廊完整配置（返回带 segment_id 别名）",
        "req": "无",
        "resp": '{"status":"success","saved":false,"operation":"queried","corridor":{corridor_id,segment_id,...}}',
    },
    {
        "name": "绿波走廊-校验配置",
        "method": "POST",
        "path": "/green_wave/validate",
        "params": "body：文档1（segment_id+green_wave_info）或文档2（corridor）结构",
        "desc": "只校验走廊配置，不落盘（dry_run）",
        "req": '文档1: {"segment_id":"green_wave_chain_1","green_wave_info":{ORDER,REF_RID,LEFT_TARGET_OFFSET_MAP,...}}\\n文档2: {"corridor":{corridor_id,intersections,periods,...}}',
        "resp": '{"status":"validated","saved":false,"message":"validation success","items":[{"segment_id":"green_wave_chain_1","saved":false,"file":{"path":"lib/green_wave_corridors.json","operation":"created"}}]}',
    },
    {
        "name": "绿波走廊-新增/更新",
        "method": "POST",
        "path": "/green_wave/update",
        "params": "body：同 validate（文档1 或文档2 结构）",
        "desc": "校验并保存走廊配置，写入 lib/green_wave_corridors.json",
        "req": "同 validate 请求结构",
        "resp": '{"status":"success","saved":true,"message":"save success","items":[{"segment_id":"green_wave_chain_1","saved":true,"file":{"path":"lib/green_wave_corridors.json","operation":"created"}}]}',
    },
    {
        "name": "绿波走廊-删除（停用）",
        "method": "POST",
        "path": "/green_wave/delete",
        "params": "body.corridor_id：走廊编号",
        "desc": "当前无彻底删除函数，转成停用（配置保留，不再参与协调）",
        "req": '{"corridor_id":"lvbo_01"}',
        "resp": '{"status":"success","saved":true,"operation":"disabled","corridor_id":"lvbo_01","enabled":false,"message":"disabled","segment_id":"lvbo_01"}',
    },
    {
        "name": "绿波走廊-启用/停用",
        "method": "PATCH",
        "path": "/green_wave/{segment_id}/enabled",
        "params": "body.enabled：true 启用 / false 停用",
        "desc": "启用或停用一条走廊，不删除配置",
        "req": '{"enabled":true}',
        "resp": '{"status":"success","saved":true,"operation":"enabled","corridor_id":"lvbo_01","enabled":true,"message":"enabled","segment_id":"lvbo_01"}',
    },
    {
        "name": "绿波走廊-彻底删除",
        "method": "DELETE",
        "path": "/green_wave/{segment_id}",
        "params": "segment_id：走廊编号",
        "desc": "彻底删除暂未实现（需 lib 新增删除函数），返回 501",
        "req": "无",
        "resp": 'HTTP 501 {"error":"彻底删除暂未实现：请改用 POST /green_wave/delete 停用，或由 lib 新增删除函数"}',
    },
]

# ---------------------------------------------------------------------------
# 渲染
# ---------------------------------------------------------------------------
def set_cn_font(run, name="微软雅黑", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cn_font(run, size={0: 20, 1: 15, 2: 12}.get(level, 11), bold=True,
                color=RGBColor(0x1F, 0x5E, 0xFF) if level == 1 else RGBColor(0x16, 0x20, 0x33))
    p.space_after = Pt(6)
    return p


def add_para(doc, text, size=10.5, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cn_font(run, size=size, bold=bold, color=color)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(4)
    return p


def build():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()
    # 默认字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(10.5)

    # 封面标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AITC 交通信号控制系统")
    set_cn_font(run, size=22, bold=True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("对外 HTTP 接口说明")
    set_cn_font(run, size=18, bold=True, color=RGBColor(0x1F, 0x5E, 0xFF))
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("版本 V1.0 · 2026-08-11 · 服务地址 {base}".format(base=BASE))
    set_cn_font(run, size=10, color=RGBColor(0x5F, 0x6B, 0x7A))
    doc.add_paragraph()

    # 通用约定
    add_heading(doc, "一、通用约定", 1)
    add_para(doc, GENERAL)

    # 接口总览
    add_heading(doc, "二、接口总览", 1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["序号", "接口", "方法", "说明"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_cn_font(run, bold=True)
    for idx, api in enumerate(INTERFACES, 1):
        row = table.add_row().cells
        values = [str(idx), api["method"] + " " + api["path"], api["method"], api["name"] + "：" + api["desc"]]
        for i, v in enumerate(values):
            row[i].text = ""
            run = row[i].paragraphs[0].add_run(v)
            set_cn_font(run, size=9.5)

    # 详细接口
    add_heading(doc, "三、接口详细说明", 1)
    for idx, api in enumerate(INTERFACES, 1):
        add_heading(doc, "{}. {}（{} {}）".format(idx, api["name"], api["method"], api["path"]), 2)
        detail = doc.add_table(rows=5, cols=2)
        detail.style = "Table Grid"
        rows = [
            ("接口路径", api["path"]),
            ("请求方法", api["method"]),
            ("请求参数", api["params"]),
            ("功能说明", api["desc"]),
        ]
        for r_i, (k, v) in enumerate(rows):
            c0 = detail.rows[r_i].cells[0]
            c1 = detail.rows[r_i].cells[1]
            c0.text = ""
            c1.text = ""
            run = c0.paragraphs[0].add_run(k)
            set_cn_font(run, bold=True, size=9.5)
            run = c1.paragraphs[0].add_run(v)
            set_cn_font(run, size=9.5)
        # 请求示例
        add_para(doc, "请求示例：", bold=True, size=9.5)
        add_code_block(doc, api["req"])
        # 响应示例
        add_para(doc, "响应示例：", bold=True, size=9.5)
        add_code_block(doc, api["resp"])
        doc.add_paragraph()

    # 附录：绿波双格式
    add_heading(doc, "四、附录：绿波接口双格式说明", 1)
    add_para(doc, "绿波走廊配置接口同时支持文档1（外方 HTTP）与文档2（Python 函数）两种请求结构，响应统一为文档1 格式：")
    add_para(doc, "1. 文档1 请求结构（segment_id + green_wave_info）：", bold=True)
    add_code_block(doc, '{\n'
                        '  "segment_id": "green_wave_chain_1",\n'
                        '  "segment_name": "绿波干线1段",\n'
                        '  "green_wave_info": {\n'
                        '    "ORDER": ["2705050","1300370","1300373","1300248"],\n'
                        '    "REF_RID": "1300370",\n'
                        '    "LEFT_OFFSET_RID": "2705050",\n'
                        '    "RIGHT_OFFSET_RID": "1300248",\n'
                        '    "LEFT_TARGET_OFFSET_MAP": {"2705050":"0","1300370":"26",...},\n'
                        '    "RIGHT_TARGET_OFFSET_MAP": {"2705050":"98","1300370":"72",...},\n'
                        '    "road_junction": {...},\n'
                        '    "topology": {...},\n'
                        '    "morning_peak_trigger": {"start":"07:30:00","end":"09:00:00","direction":"R"},\n'
                        '    "evening_peak_trigger": {"start":"17:00:00","end":"19:00:00","direction":"L"}\n'
                        '  }\n'
                        '}')
    add_para(doc, "2. 文档2 请求结构（corridor）：", bold=True)
    add_code_block(doc, '{\n'
                        '  "corridor": {\n'
                        '    "corridor_id": "lvbo_01",\n'
                        '    "name": "现有四路口绿波走廊",\n'
                        '    "enabled": true,\n'
                        '    "cycle_seconds": 90,\n'
                        '    "intersections": [{"cross_id":"2705050","green_stage_index":0,"balance_stage_index":1,"default_red_seconds":15}, ...],\n'
                        '    "periods": [{"period_id":"morning","start_time":"07:30:00","end_time":"09:00:00","reference_cross_id":"1300248","intersection_order":[...],"travel_seconds":{...}}, ...],\n'
                        '    "road_junction": {...},\n'
                        '    "topology": {...}\n'
                        '  }\n'
                        '}')
    add_para(doc, "3. 字段映射（文档1 → 文档2）：", bold=True)
    for line in [
        "segment_id → corridor_id",
        "segment_name → name",
        "green_wave_info.ORDER → intersections（阶段默认 0/1，红灯默认 0）",
        "LEFT_TARGET_OFFSET_MAP + LEFT_OFFSET_RID → direction=L 时段（evening，正向）",
        "RIGHT_TARGET_OFFSET_MAP + RIGHT_OFFSET_RID → direction=R 时段（morning，反向）",
        "morning/evening_peak_trigger.start/end → periods 起止",
        "可选顶层 enabled（默认 true）→ corridor.enabled",
    ]:
        add_para(doc, "· " + line, size=9.5)

    doc.save(OUT_FILE)
    print("文档已生成:", OUT_FILE)


if __name__ == "__main__":
    build()
